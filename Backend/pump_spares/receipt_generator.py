from docx import Document
from docx.shared import Inches
from django.http import HttpResponse
from django.conf import settings
import os
from datetime import datetime
import io

def generate_receipt(material_data, user_info=None):
    """
    Generate a receipt document with filled pump spares details
    """
    template_path = os.path.join(settings.BASE_DIR, 'Receipt.docx')
    
    if not os.path.exists(template_path):
        raise FileNotFoundError("Receipt template not found")
    
    doc = Document(template_path)
    
    current_date = datetime.now().strftime("%d/%m/%Y")
    
    if user_info and user_info.get('user_unique_code'):
        ref_no = user_info.get('user_unique_code')
    else:
        ref_no = f"SS/PS/{datetime.now().strftime('%y%m%d')}/{material_data.get('id', '001')}"
    
    if user_info:
        customer_name = user_info.get('full_name') or user_info.get('company_name') or ''
        customer_address = user_info.get('company_address') or ''
        if not customer_name:
            customer_name = ''
        if not customer_address:
            customer_address = ''
    else:
        customer_name = ''
        customer_address = ''
    
    unit_price = float(material_data.get('unit_price', 0))
    quantity = 1 
    total_price = unit_price * quantity
    
    formatted_unit_price = f"₹{unit_price:,.2f}"
    formatted_total_price = f"₹{total_price:,.2f}"
    
    replacements = {
        '[Ref. No.]': ref_no,
        '[DD/MM/YYYY]': current_date,
        '[Customer Name & Address]': f"{customer_name}\n{customer_address}",
        '[Spare Name & Details]': f"{material_data.get('part_name_value', '')} - {material_data.get('moc', '')}\n"
                                 f"Pump: {material_data.get('pump_make_name', '')} {material_data.get('pump_model_name', '')} {material_data.get('pump_size_value', '')}\n"
                                 f"Part No: {material_data.get('part_number_value', '')}",
        'X': str(quantity),
        'XXXX': formatted_unit_price,
        '[XXXX]': formatted_total_price,
        '[Location]': user_info.get('location', '') if user_info else '',
        '[X weeks]': '2-3 weeks',
        '[50%]': '50%',
        '[Your City/State]': user_info.get('location', '') if user_info else '',
        '[Your City]': user_info.get('location', '') if user_info else ''
    }
    
    for paragraph in doc.paragraphs:
        for placeholder, replacement in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, replacement in replacements.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, replacement)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def create_receipt_response(material_data, user_info=None):
    """
    Create HTTP response with the generated receipt document
    """
    try:
        buffer = generate_receipt(material_data, user_info)
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        filename = f"Pump_Spares_Receipt_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        raise Exception(f"Error generating receipt: {str(e)}")
